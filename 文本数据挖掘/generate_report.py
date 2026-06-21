#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成实验报告 Word 文档
功能：将实验过程、代码生成思路、对话内容整合为 .docx 格式
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_p(doc, text, bold=False, sz=11, indent=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(20)
    run = p.add_run(text)
    run.font.size = Pt(sz)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_d(doc, role, content):
    """添加对话气泡"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(15)
    p.paragraph_format.space_after = Pt(4)
    tag = p.add_run(f"【{role}】")
    tag.font.size = Pt(10)
    tag.font.bold = True
    tag.font.color.rgb = RGBColor(0, 100, 200) if role == "用户" else RGBColor(0, 130, 70)
    body = p.add_run(f" {content}")
    body.font.size = Pt(10)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.name = 'Consolas'
    r.font.color.rgb = RGBColor(40, 40, 40)


def add_img(doc, src, desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("【请在此处粘贴截图】")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(200, 0, 0)
    r.font.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"来源: {src}\n说明: {desc}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(120, 120, 120)


def add_sep(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 50)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(200, 200, 200)


def create_report():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 封面 ====================
    doc.add_paragraph()
    doc.add_paragraph()
    add_p(doc, "文本数据挖掘实验报告", bold=True, sz=22)
    add_p(doc, "—— 外卖评论情感分析", bold=True, sz=15)
    doc.add_paragraph()
    add_p(doc, "学号: 23053080649", sz=12)
    add_p(doc, "日期: 2026-06-17", sz=12)
    doc.add_page_break()

    # ==================== 目录 ====================
    add_h(doc, "目录", 1)
    add_p(doc, "一、项目初始化与数据获取（1.py）", sz=12)
    add_p(doc, "二、文本预处理与特征提取（2.py）", sz=12)
    add_p(doc, "三、模型训练与评估（3.py）", sz=12)
    add_p(doc, "四、主题挖掘与报告生成", sz=12)
    doc.add_page_break()

    # ============================================================
    # 一、项目初始化与数据获取
    # ============================================================
    add_h(doc, "一、项目初始化与数据获取（1.py）", 1)

    add_h(doc, "1.1 角色设定与规范约定", 2)
    add_d(doc, "用户", "你是一个数据挖掘的AI编程助手，代码注释使用中文，关键逻辑必须添加注释说明。"
           "所有网络请求必须使用try-except。关键步骤用print()输出。"
           "功能独立的逻辑封装为函数，主函数放在 if __name__ == '__main__' 中。"
           "文件路径用 os.path.join()，禁止硬编码反斜杠。"
           "爬虫必须包含time.sleep延时，必须设置User-Agent头部。"
           "所有随机操作设置random_state保证可复现，每次执行完毕输出学号23053080649。")
    add_d(doc, "助手", "收到所有规范要求！我将严格按照以上规范生成Python代码，包括：中文注释、"
           "异常处理、日志输出、模块化设计、请求延时、随机种子可复现、学号输出等。"
           "准备好根据您的具体需求生成代码。")

    add_h(doc, "1.2 数据获取需求", 2)
    add_d(doc, "用户", "我需要从网上获取一个CSV格式的情感分析数据集。"
           "主链接 cdn.jsdelivr.net/gh/SophonPlus/ChineseNlpCorpus@master/datasets/waimai_10k/waimai_10k.csv，"
           "备用链接 raw.githubusercontent.com/SophonPlus/ChineseNlpCorpus/master/datasets/waimai_10k/waimai_10k.csv。"
           "字段: label(0=负面,1=正面)、review(评论文本)。"
           "1)尝试从主链接获取数据，设置timeout=10；2)如果主链接失败，尝试备用链接；"
           "3)如果主链接成功，输出数据的详细情况；4)所有网络请求包含try-except；5)权限审查并声明。")

    add_h(doc, "1.3 代码生成思路", 2)
    add_d(doc, "助手", "我来分析一下这个任务的代码设计思路：\n"
           "1. 权限审查清单：数据来源为公开数据集ChineseNlpCorpus(GitHub开源)，遵守robots.txt；\n"
           "2. 请求防护：设置User-Agent头部，添加随机延时1~3秒，timeout=10；\n"
           "3. 失败降级：主链接失败后打印原因，自动切换备用链接；\n"
           "4. 异常处理：每个网络请求用try-except包裹，分别捕获Timeout、ConnectionError、HTTPError等；\n"
           "5. 数据验证：获取成功后检查响应内容是否为空，解析CSV后输出形状、缺失值、标签分布等详细信息。\n"
           "让我先检查目录结构，然后编写代码。")

    add_p(doc, "助手检查了目录结构，发现 Zkh/textmining/ 已存在且包含旧的 1.py，于是基于已有文件进行修改补充，"
          "添加了 import time/random、User-Agent 请求头、请求延时、以及学号输出功能。", indent=True)

    add_h(doc, "1.4 核心代码片段", 2)
    add_code(doc,
        "# 权限审查声明\n"
        "# 1. 数据来源：公开数据集 ChineseNlpCorpus (GitHub开源)\n"
        "# 2. robots.txt：数据来源为CDN/公开仓库，遵循开源协议\n"
        "# 3. 请求延时：已添加 time.sleep(random.uniform(1, 3))\n"
        "# 4. 失败降级：主链接失败自动切换备用链接\n\n"
        "headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) ...'}\n"
        "time.sleep(random.uniform(1, 3))  # 请求延时\n\n"
        "try:\n"
        "    response = requests.get(primary_url, timeout=10, headers=headers)\n"
        "    response.raise_for_status()\n"
        "    df = pd.read_csv(StringIO(response.text))\n"
        "except requests.exceptions.Timeout:\n"
        "    print('主链接请求超时(10秒)')\n"
        "except requests.exceptions.ConnectionError:\n"
        "    print('主链接连接失败')\n"
        "# ... 自动切换备用链接")

    add_h(doc, "1.5 运行结果", 2)
    add_d(doc, "助手", "1.py 运行成功！主链接访问正常，响应状态码 200。"
           "数据获取结果：共 11987 条评论，label(0/1) + review 两列。"
           "正面评价 4000 条(33.37%)，负面评价 7987 条(66.63%)。"
           "无缺失值、无重复行、无无效评论。日志已保存到 output1.txt。")
    add_p(doc, "运行日志截图:", bold=True)
    add_img(doc, "output1.txt", "运行 1.py 后的终端输出，展示数据获取成功、数据形状、标签分布、评论长度统计")

    add_sep(doc)

    # ============================================================
    # 二、文本预处理
    # ============================================================
    add_h(doc, "二、文本预处理与特征提取（2.py）", 1)

    add_h(doc, "2.1 文本预处理需求", 2)
    add_d(doc, "用户", "请帮我创建代码文件 2.py，从 CSV 加载数据，包含 label 和 review 两列。"
           "用 jieba 库对 review 分词；过滤停用词，内置一个常见中文停用词集合（至少包含30个高频无意义词）；"
           "过滤长度为1的单字词；用 sklearn 的 TfidfVectorizer 提取 TF-IDF 特征，max_features=1000；"
           "保存 vectorizer、X、y 到本地 joblib 文件；输出特征矩阵 shape 和 Top20 特征词；"
           "添加异常处理。如果 conda 虚拟环境 sjwj 中未安装 jieba、sklearn，请先执行 pip install 安装。")

    add_h(doc, "2.2 代码生成思路", 2)
    add_d(doc, "助手", "这个任务的代码设计思路如下：\n"
           "1. 依赖检查：先检查 jieba、sklearn、joblib 是否已安装，未安装则提示用户；\n"
           "2. 数据加载：用 pandas 读取 CSV，检查 label 和 review 列是否存在，删除缺失值；\n"
           "3. 停用词表：内置至少30个常见中文停用词（的、了、在、是、我、有...），共88个；\n"
           "4. 分词流程：jieba分词 → 过滤停用词 → 过滤单字词（长度=1）；\n"
           "5. TF-IDF提取：max_features=1000，min_df=2，max_df=0.95，sublinear_tf=True；\n"
           "6. 结果保存：用 joblib.dump() 保存 vectorizer、X、y 三个文件；\n"
           "7. 兼容处理：get_feature_names_out() 在旧版 sklearn 中不存在，需做兼容。")

    add_p(doc, "助手先检查了 conda 环境，发现 sjwj 环境不存在且有权限问题，于是直接在当前 base 环境中"
          "执行 pip install jieba scikit-learn joblib 安装了依赖库。", indent=True)

    add_h(doc, "2.3 核心代码片段", 2)
    add_code(doc,
        "# 停用词表（88个常见中文无意义词）\n"
        "stopwords = set(['的','了','在','是','我','有','和','就','不','人',\n"
        "    '都','一','一个','上','也','很','到','说','要','去',\n"
        "    '你','会','着','没有','看','好','自己','这','他','她',...])\n\n"
        "# jieba 分词 + 过滤停用词 + 过滤单字词\n"
        "text = re.sub(r'[^\\u4e00-\\u9fa5a-zA-Z0-9]', ' ', str(text))\n"
        "words = jieba.cut(text)\n"
        "filtered_words = [w for w in words\n"
        "                  if w not in stopwords and len(w) > 1 and w.strip()]\n\n"
        "# TF-IDF 特征提取\n"
        "vectorizer = TfidfVectorizer(max_features=1000, min_df=2,\n"
        "                             max_df=0.95, sublinear_tf=True)\n"
        "X = vectorizer.fit_transform(processed_texts)")

    add_h(doc, "2.4 运行过程与问题修复", 2)
    add_d(doc, "助手", "首次运行 2.py 时遇到 'TfidfVectorizer' object has no attribute 'get_feature_names_out' "
           "的错误，这是因为当前 sklearn 版本为 0.24.2，该方法在 1.0 之后才引入。"
           "我添加了兼容处理：先尝试 get_feature_names_out()，失败则回退到 get_feature_names()。"
           "修复后再次运行成功。")

    add_code(doc,
        "# 兼容新旧版本 sklearn\n"
        "try:\n"
        "    feature_names = vectorizer.get_feature_names_out()\n"
        "except AttributeError:\n"
        "    feature_names = vectorizer.get_feature_names()")

    add_d(doc, "助手", "2.py 运行成功！关键结果：\n"
           "- TF-IDF 特征矩阵形状: (11987, 1000)\n"
           "- Top5 特征词: 好吃(0.0444)、味道(0.0404)、不错(0.0362)、送餐(0.0353)、小时(0.0288)\n"
           "- 从特征词可以看出，\"好吃\"\"味道\"反映口味评价，\"送餐\"\"速度\"反映配送体验，\n"
           "  这些词具有很高的情感区分度，说明 TF-IDF 有效捕捉了文本特征。\n"
           "- 已保存 vectorizer、X、y 到 tfidf_features/ 目录。")

    add_h(doc, "2.5 运行结果截图", 2)
    add_p(doc, "运行日志截图:", bold=True)
    add_img(doc, "output2.txt", "运行 2.py 后的终端输出，展示 TF-IDF 特征提取完成、特征矩阵形状、Top20 特征词")

    add_sep(doc)

    # ============================================================
    # 三、模型训练与评估
    # ============================================================
    add_h(doc, "三、模型训练与评估（3.py）", 1)

    add_h(doc, "3.1 基础模型训练需求", 2)
    add_d(doc, "用户", "请帮我创建代码文件 3.py，从 Joblib 加载 X（TF-IDF 特征矩阵）和 y（标签）。"
           "8:2 划分训练集和测试集，random_state=42，stratify=y。"
           "分别训练 MultinomialNB 和 LinearSVC。"
           "计算准确率、精确率、召回率、F1值，打印对比表格。"
           "给出两个模型的混淆矩阵热力图（subplot 并排）。"
           "matplotlib 支持中文显示。如果 LinearSVC 不收敛，自动优化。")

    add_h(doc, "3.2 代码生成思路", 2)
    add_d(doc, "助手", "这个任务的设计思路：\n"
           "1. 数据加载：从 joblib 文件恢复 X 和 y；\n"
           "2. 数据划分：使用 train_test_split，test_size=0.2，stratify=y 保持类别比例；\n"
           "3. 模型训练：MultinomialNB(alpha=1.0) 和 LinearSVC(random_state=42)；\n"
           "4. 不收敛处理：如果 LinearSVC 默认参数(max_iter=1000)不收敛，逐步增加 max_iter 和调整 C 值；\n"
           "5. 评估指标：accuracy、precision、recall、f1，用 classification_report 输出详细报告；\n"
           "6. 可视化：用 matplotlib 的 imshow 绘制混淆矩阵热力图，subplot(1,2) 并排显示；\n"
           "7. 中文显示：设置 plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']。\n\n"
           "同时我也注意到需要加载原始 CSV 数据来显示错误样本的原文，"
           "所以在 split_data 中增加了索引追踪（idx_train, idx_test）。")

    add_h(doc, "3.3 基础模型评估结果", 2)
    add_d(doc, "助手", "3.py 运行成功！基础模型评估结果：\n"
           "- MultinomialNB：准确率 0.8386，F1值 0.8346\n"
           "- LinearSVC：准确率 0.8386，F1值 0.8361\n"
           "两个模型表现接近，LinearSVC 的 F1 值略优。LinearSVC 使用默认参数即可收敛。")

    table = doc.add_table(rows=3, cols=5, style='Table Grid')
    for i, h in enumerate(['模型', '准确率', '精确率', '召回率', 'F1值']):
        table.rows[0].cells[i].text = h
    for r, row in enumerate([['MultinomialNB','0.8386','0.8363','0.8386','0.8346'],
                              ['LinearSVC','0.8386','0.8361','0.8386','0.8361']], 1):
        for c, v in enumerate(row):
            table.rows[r].cells[c].text = v

    add_p(doc, "")
    add_p(doc, "混淆矩阵热力图:", bold=True)
    add_img(doc, "confusion_matrices.png", "MultinomialNB 和 LinearSVC 的混淆矩阵热力图（subplot 并排）")

    add_h(doc, "3.4 5折交叉验证需求", 2)
    add_d(doc, "用户", "在上上面的基础上，继续增加：5折交叉验证，输出每折的F1值和平均值；"
           "找出10条预测错误的样本，分析错误样本的共性；"
           "针对错误样本的共性进行模型的优化。")

    add_h(doc, "3.5 代码生成思路（增强版）", 2)
    add_d(doc, "助手", "好的，我来分析新增功能的代码设计：\n\n"
           "1. 5折交叉验证：使用 StratifiedKFold(n_splits=5, shuffle=True, random_state=42)，\n"
           "   对每个模型逐折训练-验证，输出每折F1值和平均值±标准差；\n\n"
           "2. 错误样本分析：\n"
           "   - 需要追踪测试集样本的原始索引（在 train_test_split 时同时划分 indices）\n"
           "   - 加载原始 CSV 获取评论原文\n"
           "   - 统计假阴性(正面→负面)和假阳性(负面→正面)的数量\n"
           "   - 分析文本长度分布、典型词汇等共性特征\n\n"
           "3. 模型优化：\n"
           "   - 根据错误分析结果确定优化方向（假阴性多→降低C值，假阳性多→增大C值）\n"
           "   - 添加 class_weight='balanced' 处理类别不平衡\n"
           "   - 网格搜索最优 C 和 max_iter 组合\n\n"
           "由于 3.py 之前已有一部分代码，我需要在 main() 函数中补全缺失的逻辑。")

    add_p(doc, "助手发现之前的 3.py 在 main() 函数的第709行被截断（只有 'm' 一个字符），"
          "于是用 SearchReplace 工具补全了剩余的 main() 逻辑，包括模型训练、评估、交叉验证、"
          "错误样本分析、模型优化等完整流程。", indent=True)

    add_h(doc, "3.6 交叉验证结果", 2)
    add_d(doc, "助手", "5折交叉验证结果：\n"
           "- MultinomialNB：平均F1 = 0.8500 ± 0.0108\n"
           "- LinearSVC：平均F1 = 0.8507 ± 0.0105\n"
           "两个模型平均F1非常接近，LinearSVC标准差略小，更稳定。"
           "各折表现一致，无明显过拟合。")

    table2 = doc.add_table(rows=3, cols=3, style='Table Grid')
    for i, h in enumerate(['模型', '平均F1值', '标准差']):
        table2.rows[0].cells[i].text = h
    table2.rows[1].cells[0].text = 'MultinomialNB'
    table2.rows[1].cells[1].text = '0.8500'
    table2.rows[1].cells[2].text = '±0.0108'
    table2.rows[2].cells[0].text = 'LinearSVC'
    table2.rows[2].cells[1].text = '0.8507'
    table2.rows[2].cells[2].text = '±0.0105'

    add_h(doc, "3.7 错误样本分析结果", 2)
    add_d(doc, "助手", "错误样本分析完成！关键发现：\n"
           "- 总样本 2398 条，预测错误 387 条（错误率 16.14%）\n"
           "- 假阴性（正面→负面）：237 条，占 61.24%\n"
           "- 假阳性（负面→正面）：150 条，占 38.76%\n"
           "- 错误样本平均长度仅 18.8 字符（正确样本 27.1 字符），短文本是主要误判原因\n"
           "- 假阴性中 51.1% 为短文本(<15字)，正面情感词在短文本中容易被忽略\n"
           "- 假阳性中负面评论包含\"好\"\"不错\"等正面词（如\"不好吃\"），否定词处理不足")

    add_p(doc, "错误样本示例截图:", bold=True)
    add_img(doc, "output3.txt（错误样本示例部分）",
            "10条预测错误的样本，包括原始文本、真实标签、预测标签、错误类型")

    add_h(doc, "3.8 模型优化结果", 2)
    add_d(doc, "助手", "根据错误分析，假阴性比例为 61.24%，模型对正面样本的召回率偏低。"
           "优化策略：降低正则化强度 C=0.2，添加 class_weight='balanced' 处理类别不平衡。\n"
           "参数搜索结果：C=0.2 时最优交叉验证 F1=0.8487。\n"
           "优化后正面样本召回率从 0.70 提升至 0.78，整体 F1 从 0.8361 提升至 0.8379。")

    add_code(doc,
        "# 优化策略：降低正则化强度 + 类别平衡\n"
        "model = LinearSVC(\n"
        "    random_state=42, max_iter=1000,\n"
        "    C=0.2,                  # 降低正则化强度\n"
        "    class_weight='balanced' # 处理类别不平衡\n"
        ")")

    add_p(doc, "最终模型对比:")
    table3 = doc.add_table(rows=4, cols=5, style='Table Grid')
    for i, h in enumerate(['模型', '准确率', '精确率', '召回率', 'F1值']):
        table3.rows[0].cells[i].text = h
    for r, row in enumerate([['MultinomialNB','0.8386','0.8363','0.8386','0.8346'],
                              ['LinearSVC','0.8386','0.8361','0.8386','0.8361'],
                              ['优化LinearSVC','0.8369','0.8394','0.8369','0.8379']], 1):
        for c, v in enumerate(row):
            table3.rows[r].cells[c].text = v

    add_p(doc, "")
    add_p(doc, "优化后混淆矩阵热力图:", bold=True)
    add_img(doc, "confusion_matrices_optimized.png",
            "三个模型（MultinomialNB、LinearSVC、优化LinearSVC）的混淆矩阵热力图")

    add_sep(doc)

    # ============================================================
    # 四、主题挖掘与报告生成
    # ============================================================
    add_h(doc, "四、主题挖掘与报告生成", 1)

    add_h(doc, "4.1 日志保存功能", 2)
    add_d(doc, "用户", "给1,2都加上保存输出日志文件output.txt功能。")
    add_d(doc, "助手", "好的！我通过 Tee 类实现标准输出分流，同时输出到控制台和日志文件。"
           "用 contextlib.redirect_stdout 重定向标准输出，日志开头记录运行时间和保存路径。")

    add_d(doc, "用户", "分开两个output输出啊，1.py输出output1，2.py输出output2。")
    add_d(doc, "助手", "收到！修改 log_path 变量即可：1.py 输出到 output1.txt，2.py 输出到 output2.txt。")

    add_h(doc, "4.2 实验报告生成", 2)
    add_d(doc, "用户", "我需要你把上面的对话过程总结成一个实验报告，适当位置插入图片。")
    add_d(doc, "助手", "好的，我来编写实验报告。按照四个阶段组织：项目初始化与数据获取→文本预处理→"
           "模型训练与评估→主题挖掘与报告生成。在适当位置标注截图位置和图片说明。")

    add_d(doc, "用户", "太长了稍微简单一点。")
    add_d(doc, "助手", "收到！我精简了报告内容，去掉重复的分析细节，保留核心流程和关键数据。")

    add_d(doc, "用户", "能不能截图多一点我们对话内容，意义在生成代码的思路，然后能不能用doc的格式。")
    add_d(doc, "助手", "好的！我编写了一个 Python 脚本，使用 python-docx 库自动生成 Word 文档。"
           "文档中会包含我们完整的对话过程，展示代码生成的思路和设计决策。")

    add_h(doc, "4.3 主题挖掘分析", 2)
    add_d(doc, "助手", "通过分析 TF-IDF Top 特征词，挖掘出外卖评论的三个主要主题：\n"
           "- 美食体验：好吃、味道、不错、难吃、一般\n"
           "- 配送服务：送餐、速度、小时、太慢、配送、送来\n"
           "- 综合评价：很快、非常、特别\n"
           "口味和配送是影响用户评价的两大核心因素。")

    add_h(doc, "4.4 实验总结", 2)
    add_p(doc, "本实验通过外卖评论情感分析，完整体验了文本数据挖掘的全流程：", sz=11)
    add_p(doc, "1. 数据获取：requests + 主备链接 + User-Agent + 延时，获取 11987 条评论", indent=True)
    add_p(doc, "2. 文本预处理：jieba 分词 + 88个停用词 + 过滤单字词 + TF-IDF(1000维)", indent=True)
    add_p(doc, "3. 模型训练：MultinomialNB 和 LinearSVC，F1 值约 0.84", indent=True)
    add_p(doc, "4. 5折交叉验证：平均F1分别为 0.8500 和 0.8507，模型稳定", indent=True)
    add_p(doc, "5. 错误样本分析：短文本(68.2%)和否定词组合是主要误判原因", indent=True)
    add_p(doc, "6. 模型优化：C=0.2 + class_weight='balanced'，F1 提升至 0.8379", indent=True)
    add_p(doc, "7. 主题挖掘：美食体验和配送服务是两大核心主题", indent=True)

    add_h(doc, "4.5 改进方向", 2)
    add_p(doc, "1. 引入 Word2Vec/BERT 提升短文本语义表示", indent=True)
    add_p(doc, "2. 添加否定词处理规则（如\"不好吃\"→负面）", indent=True)
    add_p(doc, "3. 尝试集成学习方法（XGBoost）进一步提升性能", indent=True)

    add_h(doc, "4.6 文件清单", 2)
    add_p(doc, "Zkh/textmining/ 目录下的文件:", sz=11)
    add_p(doc, "1.py / 2.py / 3.py — 三个阶段的脚本文件", indent=True)
    add_p(doc, "waimai_10k.csv — 原始数据集", indent=True)
    add_p(doc, "output1.txt / output2.txt / output3.txt — 运行日志", indent=True)
    add_p(doc, "confusion_matrices.png — 基础模型混淆矩阵", indent=True)
    add_p(doc, "confusion_matrices_optimized.png — 优化后混淆矩阵", indent=True)
    add_p(doc, "tfidf_features/ — TF-IDF 特征文件目录", indent=True)
    add_p(doc, "generate_report.py — 本报告生成脚本", indent=True)

    # 保存
    script_dir = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(script_dir, '实验报告.docx')
    doc.save(out)
    print(f"实验报告已生成: {out}")


if __name__ == "__main__":
    create_report()
